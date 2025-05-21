import os
import re
import glob
from dotenv import load_dotenv
from docx import Document
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS

# =================== LOAD API KEY ===================
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("❌ OPENAI_API_KEY tidak ditemukan di environment.")

# =================== MUAT PDF ===================
print("🚀 Memuat dokumen PDF dari folder 'data/'...")
pdf_files = glob.glob("data/*.pdf")
print(f"📂 Total file PDF ditemukan: {len(pdf_files)}")

loader = DirectoryLoader('data', glob="./*.pdf", loader_cls=PyPDFLoader)
documents = loader.load()
print(f"📄 Total halaman (dokumen) yang dimuat: {len(documents)}")

# =================== SPLIT DOKUMEN ===================
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
texts = text_splitter.split_documents(documents)
print(f"🧩 Total chunk yang dihasilkan: {len(texts)}")

# =================== HITUNG HURUF ===================
def count_letters(text_chunks):
    total = 0
    for chunk in text_chunks:
        clean = re.sub(r'[^a-zA-Z]', '', chunk.page_content)
        total += len(clean)
    return total

total_letters = count_letters(texts)
print(f"🔤 Total huruf: {total_letters}")

# =================== SIMPAN KE DOCX ===================
doc = Document()
doc.add_heading("Isi Chunk Dokumen", 0)
for idx, chunk in enumerate(texts, start=1):
    doc.add_heading(f"Chunk {idx}", level=2)
    doc.add_paragraph(chunk.page_content.strip())
doc.save("log_karakter.docx")
print("📄 log_karakter.docx disimpan.")

# =================== EMBEDDING & SIMPAN FAISS ===================
print("🔄 Membuat embeddings (OpenAI - text-embedding-ada-002) dan menyimpan FAISS...")

embeddings = OpenAIEmbeddings(
    model="text-embedding-ada-002",
    openai_api_key=OPENAI_API_KEY
)

faiss_db = FAISS.from_documents(texts, embeddings)
faiss_db.save_local("vectorstores/matkul_db_faiss")

print("✅ Selesai. FAISS disimpan ke vectorstores/matkul_db_faiss")
